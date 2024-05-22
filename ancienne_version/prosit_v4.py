import docx
from docx.shared import Inches
import os
import openai
import re

def generate_text(prompt):
    
    openai.api_key = "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"

    completions = openai.Completion.create(
        engine="text-davinci-002",
        prompt=prompt,
        max_tokens=1024,
        n=1,
        stop=None,
        temperature=0.5,
    )

    message = completions.choices[0].text
    return message.strip()

tab_file = []
files = os.listdir('.')
for name in files:
    tab_file.append(name)

tab_file.remove('logo_cesi.jpg')
tab_file.remove('prosit_v4.py')
tab_file.remove('sauvegarde')
doc_name = tab_file[0]

print("transformation de : " + doc_name)

match = re.search(r'\d+', doc_name)
if match:
	var_num = match.group()
print(var_num)

doc_original = docx.Document(doc_name)
size_doc = len(doc_original.paragraphs)

tab_categorie = ['Analyse du contexte','Contexte','analyse du contexte ','Mots clés', 'Mots-clés','Problématique','Problématiques','Contraintes', 'Contrainte', 'Livrables', 'Livrable','Généralisation','Généralisation',
	'Piste de solutions','Pistes de solution','Piste de solution','Pistes de solutions','Plans d’action','Plan d’actions','Plans d’actions', 'Plan d’action','Réalisation du plan d’action ']

doc_final = docx.Document()

doc_final.add_picture("logo_cesi.jpg", width=Inches(2.5))
a = doc_final.add_heading(doc_original.paragraphs[0].text, level=1)
a.alignment = 1

var_cat_mc = False

var_cat_mc = False

for i in range (1, size_doc):
	t = True
	for j in range(0, len(tab_categorie)):

		if doc_original.paragraphs[i].text.strip().lower() == tab_categorie[j].lower():
			t = False

			if doc_original.paragraphs[i].text.strip().lower() == 'Mots-clés'.lower() or doc_original.paragraphs[i].text.strip().lower() == 'Mots clés'.lower():
				print("Mots-clés")
				doc_final.add_heading(tab_categorie[j], level=2)
				var_cat_mc = True

			else:
				print("titre paragraphe")
				doc_final.add_heading(tab_categorie[j], level=2)
				var_cat_mc = False

	if t == True and var_cat_mc == False:
		print("paragraphe")
		u = doc_final.add_paragraph(doc_original.paragraphs[i].text)
		u.paragraph_format.line_spacing = 1 # interligne à 0

	if t == True and var_cat_mc == True:
		print("Mots-clés a definir")
		u = doc_final.add_paragraph(generate_text("Quelle est la défintion en BTP de " + doc_original.paragraphs[i].text))
		u.paragraph_format.line_spacing = 1 # interligne à 0


doc_fini_name = "CER_prosit_" + var_num + "_hugo_laplace.docx"
doc_final.save(doc_fini_name)

print("fichier editer avec succes")
os.system(doc_fini_name) # ouverture fichier fini

#https://stackoverflow.com/questions/24031011/python-docx-library-text-align
#https://python-docx.readthedocs.io/en/latest/